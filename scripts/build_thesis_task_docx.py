# 生成《毕业论文（设计）任务书》Word，需: pip install python-docx
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm


def ensure_run_eastasia(run, eastasia: str):
    run.font.name = eastasia
    r = run._element
    if r.rPr is None:
        r.get_or_add_rPr()
    rPr = r.rPr
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), eastasia)


def main():
    # 使用 ASCII 文件名避免部分环境下路径乱码；可在资源管理器中重命名为中文名
    out = Path(__file__).resolve().parents[1] / "thesis_task_book_drowning_monitoring.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("毕业论文（设计）任务书")
    r.bold = True
    r.font.size = Pt(22)
    ensure_run_eastasia(r, "黑体")

    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("课题名称", "基于YOLO与深度学习的泳池溺水智能监测与预警系统"),
        ("姓    名", "宋宇"),
        ("学    号", "202206040231"),
        ("院    系", "信息工程学院"),
        ("专    业", "数据科学与大数据技术"),
        ("指导教师", "周振宇 助教"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        for c in table.rows[i].cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(14)
                    ensure_run_eastasia(run, "楷体_GB2312" if i > 0 else "华文中宋")
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        ensure_run_eastasia(table.rows[i].cells[0].paragraphs[0].runs[0], "华文中宋")

    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = d.add_run("2026 年 3 月 31 日")
    dr.font.size = Pt(14)
    ensure_run_eastasia(dr, "华文中宋")

    doc.add_paragraph()

    outer = doc.add_table(rows=1, cols=1)
    outer.style = "Table Grid"
    cell = outer.cell(0, 0)

    def add_para(text, bold=False, first_indent=True):
        p = cell.add_paragraph()
        if first_indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        ensure_run_eastasia(run, "宋体")
        if bold:
            run.bold = True
        return p

    add_para(
        "关键词：溺水监测；YOLOv8；行为识别；深度学习；智能预警",
        bold=True,
        first_indent=False,
    )

    body = (
        "本课题实现一套「桌面实时监测 + 本地日志 + 可选远程通知」的溺水智能预警原型：在单机上用摄像头或视频文件持续分析画面，"
        "先用人检测框定泳客，再对裁剪区域做溺水/游泳/离水三类判别，并用连续多帧结果综合判定是否触发报警，避免单帧误判。"
        "报警时除界面高亮与系统托盘提示外，自动留存截图与短时录像，事件以 CSV 与 JSONL 形式落地，便于事后追溯；"
        "另起独立进程提供只读 Web 看板，仅在局域网内汇总展示，与检测主程序互不阻塞。"
        "远程侧支持钉钉群机器人加签推送；简报类功能仅调用 DeepSeek 等接口处理文本日志，不上传原始视频。"
        "配置与密钥放在本地 settings.json，不进入代码仓库，便于在泳池试部署时单独交接。"
    )
    add_para(body)

    add_para("须具备以下功能：", bold=True, first_indent=True)

    items = [
        "视频接入与监测参数：支持 USB 摄像头与本地视频文件；可划定多边形水域监控区，未划定时默认全画面分析。可调检测灵敏度、溺水状态连续确认帧数、报警冷却时间，随时启停分析。",
        "人体检测与目标跟踪：采用主流一阶段目标检测模型框出泳客，按框裁出人体图像供分类；支持多目标同时出现时的跟踪与编号，便于多帧累计判定。",
        "溺水相关状态识别：将裁图分为溺水、正常游泳、离水三类；分类模型可由用户按自建数据集训练或更换权重，与检测模型分工协作。",
        "多帧确认与报警：仅在连续若干帧均倾向溺水时才认定报警，并设冷却间隔避免同一场景反复刷屏；界面在确认前可呈现高风险提示与确认后告警两种层次。",
        "现场存证与事件记录：确认报警时自动保存关键帧截图与一段短时录像，并写入带时间戳、现场人数、置信度及媒体路径的结构化记录，供事后调取与统计。",
        "远程消息推送：支持向钉钉等工作群组推送文字卡片式告警，手机端可即时查看，与现场声光/桌面提示形成互补。",
        "独立监测看板（可选）：在局域网内用浏览器访问只读页面，汇总告警次数与近期事件列表；与实时监测分属不同进程，监测端不依赖浏览器是否打开。",
        "智能简报（可选）：在填写大模型服务密钥后，可基于近期事件文本生成一段话述摘要；未配置时核心监测与日志仍完整可用，且不向公网上传原始视频流。",
        "桌面客户端：单窗口集成实时画面、参数滑块与状态提示，可最小化至托盘并配合操作系统通知，便于机房长期值守。",
        "数据与实验：准备训练集与验证集，完成分类模型训练，给出准确率等指标及简单对比实验，支撑论文中的方法与结果章节。",
    ]
    for i, it in enumerate(items, 1):
        add_para(f"{i}）{it}")

    add_para("具体要求如下：", bold=True)

    reqs = [
        "开题与综述：完成文献综述与技术路线说明，交代本课题的定位与适用场景，并对数据与隐私问题有基本讨论。",
        "设计与实现：完成系统总体设计与实现，说明主要模块关系；完成模型训练与实验，给出评价指标与简要分析；说明系统在常见运行环境下的表现。",
        "成果与答辩：提交可运行的工程与说明文档，完成论文与答辩材料，现场或录屏演示系统功能，能说明系统局限与改进思路。",
    ]
    for i, it in enumerate(reqs, 1):
        add_para(f"{i}）{it}")

    add_para("文献查阅指引：", bold=True)

    # 中文著录为主；[8]–[10] 为方法类原文，英文著录便于与 CVPR/NeurIPS/Springer 对照
    refs = [
        "雷飞, 朱恒宇, 唐菲菲, 王鑫源. 基于深度学习的游泳池溺水行为检测[J]. Signal, Image and Video Processing, 2022, 16(6): 1683-1690. DOI:10.1007/s11760-021-02124-9.",
        "王帆, 艾洋, 张伟东. 基于视频监控的室内游泳池深水区早期危险状态检测[J]. Signal, Image and Video Processing, 2022, 16(1): 29-37. DOI:10.1007/s11760-021-01953-y.",
        "蒋欣航, 唐多勋, 徐文申, 张颖, 林烨. 多游泳场景下基于改进 YOLO 算法的溺水检测方法: Swimming-YOLO[J]. Signal, Image and Video Processing, 2025, 19(2). DOI:10.1007/s11760-024-03744-7.",
        "宋奇, 姚博丹, 薛云龙, 季树德. 面向溺水检测的轻量高精度 MS-YOLO 模型[J]. Sensors, 2024, 24(21): 6955. DOI:10.3390/s24216955.",
        "杨瑞亮, 王开凯, 杨立斌. 室内游泳池溺水检测的改进 YOLOv5 算法[J]. Applied Sciences, 2024, 14(1): 200. DOI:10.3390/app14010200.",
        "何天仪, 叶晓东, 王美玲. 基于 YOLOv8 的改进游泳池溺水检测方法[C]//2023 IEEE 第7届信息技术与机电工程会议 (ITOEC). Piscataway: IEEE, 2023: 835-839. DOI:10.1109/itoec57671.2023.10291322.",
        "何昕宇, 苑菲, 刘廷壮, 朱艺. 基于卷积自编码器的溺水检测视频系统[J]. Neural Computing and Applications, 2023, 35(21): 15791-15803. DOI:10.1007/s00521-023-08526-9.",
        "Kukartsev V V, Ageev R A, Borodulin A S, et al. Deep learning for object detection in images: development and evaluation of the YOLOv8 model using Ultralytics and Roboflow libraries[M]//Software Engineering Methods: Design and Application. Cham: Springer Nature Switzerland, 2024: 629-637. DOI:10.1007/978-3-031-70285-3_48.",
        "何恺明, 张祥雨, 任少卿, 孙剑. 深度残差学习用于图像识别[C]//IEEE Conference on Computer Vision and Pattern Recognition (CVPR). Las Vegas: IEEE, 2016: 770-778.",
        "Paszke A, Gross S, Massa F, et al. PyTorch: an imperative style, high-performance deep learning library[C]//Advances in Neural Information Processing Systems (NeurIPS 2019). Vancouver: Curran Associates Inc., 2019: 8024-8035.",
    ]
    for i, ref in enumerate(refs, 1):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{i}] {ref}")
        r.font.size = Pt(10.5)
        ensure_run_eastasia(r, "宋体")

    add_para("4、毕业论文（设计）进度安排：", bold=True, first_indent=False)

    # 共 21 周、四阶段（5+7+6+3）
    sched = [
        "第一阶段（第1～5周）：文献调研与开题；明确「辅助预警、不替代救生员」定位及三分类与连续帧策略；搭建 Python/PyTorch 环境，整理 data_cls 与标注数据。",
        "第二阶段（第6～12周）：跑通 prepare_dataset、train_classifier，得到分类权重与训练曲线；完成检测数据与基线指标记录。",
        "第三阶段（第13～18周）：串联 VideoProcessor，调 ROI、确认帧与录像；打通 events、FastAPI 看板与钉钉；试跑并记录误报/漏报。",
        "第四阶段（第19～21周）：论文初稿与修改、查重与格式；演示视频与答辩材料；按学院要求提交终稿并答辩。",
    ]
    for s in sched:
        p = cell.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(s)
        r.font.size = Pt(10.5)
        ensure_run_eastasia(r, "宋体")

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("院（系）意见：\n\n\n负责人签名：________________    ")
    r.font.size = Pt(10.5)
    ensure_run_eastasia(r, "宋体")

    doc.save(out)
    print("已生成:", out)


if __name__ == "__main__":
    main()
